from __future__ import annotations

import io
from datetime import datetime

import structlog
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from starlette.responses import StreamingResponse

from takehome.db.models import Message
from takehome.db.session import get_session
from takehome.services.conversation import get_conversation

logger = structlog.get_logger()

router = APIRouter(tags=["export"])


def _safe_filename(title: str) -> str:
    safe = "".join(c if c.isalnum() or c in " -_" else "_" for c in title)
    return safe.strip().replace(" ", "_") or "conversation"


def _add_paragraph(
    doc: DocxDocument,
    text: str,
    size: int,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:  # type: ignore[type-arg]
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _build_docx(
    title: str,
    documents: list[tuple[str, int]],
    messages: list[Message],
    exported_at: datetime,
) -> io.BytesIO:
    doc = DocxDocument()

    # Title
    doc.add_heading(title, level=1)

    # Export timestamp
    _add_paragraph(
        doc,
        f"Exported {exported_at.strftime('%d %B %Y at %H:%M')}",
        size=9,
        italic=True,
        color=RGBColor(0x9C, 0x9C, 0x9C),
    )

    # Documents reviewed
    if documents:
        doc.add_paragraph()
        _add_paragraph(
            doc, "Documents Reviewed", size=11, bold=True, color=RGBColor(0x44, 0x44, 0x44)
        )
        for filename, page_count in documents:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(filename)
            suffix = p.add_run(f"  —  {page_count} page{'s' if page_count != 1 else ''}")
            suffix.font.size = Pt(9)
            suffix.font.color.rgb = RGBColor(0x9C, 0x9C, 0x9C)

    # Divider
    doc.add_paragraph()
    _add_paragraph(doc, "─" * 60, size=9, color=RGBColor(0xCC, 0xCC, 0xCC))
    doc.add_paragraph()

    # Messages
    for msg in messages:
        if msg.role not in ("user", "assistant"):
            continue

        if msg.role == "user":
            _add_paragraph(doc, "You", size=10, bold=True, color=RGBColor(0x44, 0x44, 0x44))
        else:
            _add_paragraph(
                doc, "AI Assistant", size=10, bold=True, color=RGBColor(0x0D, 0x6E, 0x4E)
            )

        # Body — add_paragraph with text directly; safe even for empty strings
        body = doc.add_paragraph()
        body_run = body.add_run(msg.content)
        body_run.font.size = Pt(11)

        if msg.role == "assistant" and msg.sources_cited > 0:
            _add_paragraph(
                doc,
                f"{msg.sources_cited} source{'s' if msg.sources_cited != 1 else ''} cited",
                size=9,
                italic=True,
                color=RGBColor(0x9C, 0x9C, 0x9C),
            )

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@router.get("/api/conversations/{conversation_id}/export")
async def export_conversation(
    conversation_id: str,
    session: AsyncSession = Depends(get_session),
) -> StreamingResponse:
    """Export a conversation as a Word (.docx) document."""
    conversation = await get_conversation(session, conversation_id)
    if conversation is None:
        raise HTTPException(status_code=404, detail="Conversation not found")

    stmt = (
        select(Message)
        .where(Message.conversation_id == conversation_id)
        .order_by(Message.created_at.asc())
    )
    result = await session.execute(stmt)
    messages = list(result.scalars().all())

    if not messages:
        raise HTTPException(status_code=400, detail="No messages to export")

    documents = [(d.filename, d.page_count) for d in conversation.documents]

    buffer = _build_docx(
        title=conversation.title,
        documents=documents,
        messages=messages,
        exported_at=datetime.now(),
    )

    filename = f"{_safe_filename(conversation.title)}.docx"
    logger.info("Conversation exported", conversation_id=conversation_id, filename=filename)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
